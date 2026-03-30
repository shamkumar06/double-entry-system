from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sqlalchemy.orm import Session
from app.models.orm import Project, Phase, Category
import os
import tempfile
import logging
from collections import defaultdict
from datetime import datetime

from app.services.accounting import get_journal, get_trial_balance, get_ledger_page

logger = logging.getLogger(__name__)

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def apply_studio_table_styles(table, is_header=False):
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(6)
    
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, "F2F2F2")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def generate_project_word_report(
    db: Session,
    project_id: int, 
    report_type: str = "Journal", 
    phase_id: str = None, 
    account_name: str = "",
    date_format: str = "YYYY-MM-DD",
    sort_order: str = "Descending",
    sections: dict[str, bool] = None,
    custom_header: str = "",
    sub_headers: list[dict] = [],
    footer_note: str = "",
    show_date_corner: bool = True,
    columns: dict[str, list[str]] = None,
    start_date: str = None,
    end_date: str = None,
    ledger_accounts: list[str] = [],
    header_font_size: int = 26
) -> str:
    """
    Generates a Word Document financial report using the SQL database.
    """
    # 1. Resolve Project Data
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise Exception(f"Project {project_id} not found")
    
    project_name = project.name

    def format_dt(dt_str: str):
        if not dt_str: return ""
        d = dt_str.split('T')[0]
        if date_format == "DD/MM/YYYY":
            parts = d.split('-')
            if len(parts) == 3: return f"{parts[2]}/{parts[1]}/{parts[0]}"
        return d

    def format_curr(amt: float):
        s = f"{float(amt or 0):,.2f}"
        if s.endswith(".00"): s = s[:-3]
        return f"Rs. {s}"

    # 2. Fetch Data
    journal_data = get_journal(db, project_id, phase_id=phase_id)
    if start_date: journal_data = [tx for tx in journal_data if tx.get("transaction_date", "") >= start_date]
    if end_date: journal_data = [tx for tx in journal_data if tx.get("transaction_date", "") <= end_date]
    
    trial_balance_data = get_trial_balance(db, project_id, phase_id=phase_id)

    # 3. Document Setup
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'; font.size = Pt(11)
    
    # 4. Branding & Header
    if show_date_corner:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Report Date: {datetime.now().strftime('%d/%m/%Y')}")
        run.font.size = Pt(9); run.font.color.rgb = RGBColor(100, 100, 100)
    
    main_title = custom_header if custom_header else project_name
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(main_title.upper())
    title_run.bold = True
    title_run.font.size = Pt(header_font_size)
    
    p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("_" * 50); run.bold = True

    for sub in sub_headers:
        text = sub.get("text", "")
        if text and text.strip():
            p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.strip())
            run.font.size = Pt(sub.get("font_size", 12))
            run.font.color.rgb = RGBColor(50, 50, 50)
    
    if phase_id:
        phase = db.query(Phase).filter(Phase.id == phase_id, Phase.project_id == project_id).first()
        target_phase_name = phase.name if phase else "Selected Phase"
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Phase: {target_phase_name}")
        run.font.size = Pt(10); run.italic = True

    is_full_report = report_type in ["Full", "Dashboard", "Double Entry"]

    # JOURNAL SECTION
    show_journal = (is_full_report or report_type == "Journal")
    if sections and is_full_report:
        show_journal = sections.get("journal", True)

    if show_journal:
        document.add_paragraph()
        h = document.add_paragraph()
        run = h.add_run('JOURNAL ENTRIES'); run.bold = True; run.font.size = Pt(16)

        if not journal_data:
            document.add_paragraph("No transactions found.")
        else:
            rev = (sort_order == "Descending")
            journal_data.sort(key=lambda x: x.get('transaction_date', ''), reverse=rev)

            available_cols = {"Date": "transaction_date", "Phase": "phase_name", "From": "from_name", "To": "to_name", "Category": "category_name", "Description": "description", "Amount": "amount"}
            active_cols = columns.get("journal") if (columns and columns.get("journal")) else list(available_cols.keys())
            
            table = document.add_table(rows=1, cols=len(active_cols))
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, col_name in enumerate(active_cols):
                hdr[i].text = col_name
                hdr[i].paragraphs[0].runs[0].bold = True
                hdr[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            
            widths = [Inches(0.85), Inches(0.75), Inches(0.85), Inches(0.85), Inches(0.85), Inches(1.42), Inches(1.1)]
            set_column_widths(table, widths[:len(active_cols)])
                
            for tx in journal_data:
                row = table.add_row().cells
                for i, col_name in enumerate(active_cols):
                    key = available_cols.get(col_name); val = tx.get(key, "")
                    if col_name == "Date": row[i].text = format_dt(str(val))
                    elif col_name == "Amount": row[i].text = format_curr(val)
                    elif col_name == "Phase": row[i].text = str(val or "Whole")
                    else: row[i].text = str(val or "")
                    row[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            apply_studio_table_styles(table)

    # LEDGER SECTION
    show_ledger = (is_full_report or report_type == "Ledger")
    if sections and is_full_report:
        show_ledger = sections.get("ledger", True)

    if show_ledger:
        document.add_paragraph()
        h = document.add_paragraph()
        run = h.add_run('GENERAL LEDGER'); run.bold = True; run.font.size = Pt(16)
        
        categories = db.query(Category).all()
        cat_list = [{"name": c.name, "code": c.code} for c in categories]
        cat_list = [{"name": "Cash", "code": 1001}, {"name": "Bank", "code": 1002}] + cat_list

        ledger_targets = ledger_accounts if ledger_accounts else [c["name"] for c in cat_list]

        found_any = False
        for target_name in ledger_targets:
            cat_obj = next((c for c in cat_list if c["name"] == target_name), None)
            if not cat_obj: continue
            
            entries = get_ledger_page(db, project_id, cat_obj["code"], phase_id=phase_id)
            if not entries: continue
            found_any = True

            if sort_order == "Descending": entries.reverse()

            p = document.add_paragraph()
            run = p.add_run(f'ACCOUNT: {target_name.upper()} (Code: {cat_obj["code"]})')
            run.bold = True; run.font.size = Pt(12)
            
            active_cols = columns.get("ledger") if (columns and columns.get("ledger")) else ['Date', 'Phase', 'Debit', 'Credit', 'Running Balance']
            table = document.add_table(rows=1, cols=len(active_cols))
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, c_name in enumerate(active_cols):
                hdr[i].text = c_name
                hdr[i].paragraphs[0].runs[0].bold = True
                hdr[i].paragraphs[0].runs[0].font.size = Pt(9.5)

            set_column_widths(table, [Inches(1.2)] * len(active_cols))
            
            for entry in entries:
                row = table.add_row().cells
                for i, c_name in enumerate(active_cols):
                    if c_name == 'Date': row[i].text = format_dt(entry["date"])
                    elif c_name == 'Phase': row[i].text = str(entry.get("phase_name") or "Whole")
                    elif c_name == 'Debit': row[i].text = format_curr(entry['amount']) if entry["entry_type"] == "Debit" else "-"
                    elif c_name == 'Credit': row[i].text = format_curr(entry['amount']) if entry["entry_type"] == "Credit" else "-"
                    elif c_name == 'Running Balance': row[i].text = format_curr(entry.get('running_balance', 0))
                    row[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            apply_studio_table_styles(table)
        
        if not found_any:
            document.add_paragraph("No ledger entries found for selected filters.")

    # TRIAL BALANCE SECTION
    show_tb = (is_full_report or report_type == "Trial Balance")
    if sections and is_full_report:
        show_tb = sections.get("trialBalance", True)

    if show_tb:
        document.add_paragraph()
        h = document.add_paragraph()
        run = h.add_run('TRIAL BALANCE'); run.bold = True; run.font.size = Pt(16)
        
        accounts_tb = trial_balance_data.get("accounts", {})
        if not accounts_tb:
            document.add_paragraph("No balances found.")
        else:
            table = document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Account Name', 'Debit Balance', 'Credit Balance'
            for cell in hdr: 
                cell.paragraphs[0].runs[0].bold = True; cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            set_column_widths(table, [Inches(3.27), Inches(1.7), Inches(1.7)])
            
            for acc, adata in accounts_tb.items():
                row = table.add_row().cells
                row[0].text = acc
                bal = adata.get("balance", 0.0)
                if bal >= 0:
                    row[1].text = format_curr(bal); row[2].text = "-"
                else:
                    row[1].text = "-"; row[2].text = format_curr(abs(bal))
                for cell in row: cell.paragraphs[0].runs[0].font.size = Pt(10)
                
            footer = table.add_row().cells
            footer[0].text = "TOTAL"; footer[0].paragraphs[0].runs[0].bold = True
            footer[1].text = format_curr(trial_balance_data['totals']['total_debits']); footer[1].paragraphs[0].runs[0].bold = True
            footer[2].text = format_curr(trial_balance_data['totals']['total_credits']); footer[2].paragraphs[0].runs[0].bold = True
            apply_studio_table_styles(table)

    if footer_note:
        document.add_paragraph()
        note_text = document.add_paragraph()
        note_run = note_text.add_run(footer_note)
        note_run.italic = True; note_run.font.size = Pt(10); note_run.font.color.rgb = RGBColor(100, 100, 100)

    temp_dir = tempfile.gettempdir()
    safe_name = "".join(x for x in project_name if x.isalnum())
    file_path = os.path.join(temp_dir, f"{safe_name}_{report_type.replace(' ', '_')}.docx")
    document.save(file_path); return file_path

    is_full_report = report_type in ["Full", "Dashboard", "Double Entry"]

    # JOURNAL SECTION
    show_journal = (is_full_report or report_type == "Journal")
    if sections and is_full_report:
        show_journal = sections.get("journal", True)

    if show_journal:
        document.add_paragraph()
        h = document.add_paragraph()
        run = h.add_run('JOURNAL ENTRIES'); run.bold = True; run.font.size = Pt(16)

        if not journal_data:
            document.add_paragraph("No transactions found.")
        else:
            rev = (sort_order == "Descending")
            journal_data.sort(key=lambda x: x.get('transaction_date', ''), reverse=rev)

            available_cols = {"Date": "transaction_date", "Phase": "phase_name", "From": "from_name", "To": "to_name", "Category": "category_name", "Description": "description", "Amount": "amount"}
            active_cols = columns.get("journal") if (columns and columns.get("journal")) else list(available_cols.keys())
            
            table = document.add_table(rows=1, cols=len(active_cols))
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, col_name in enumerate(active_cols):
                hdr[i].text = col_name
                hdr[i].paragraphs[0].runs[0].bold = True
                hdr[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            
            widths = [Inches(0.85), Inches(0.75), Inches(0.85), Inches(0.85), Inches(0.85), Inches(1.42), Inches(1.1)]
            set_column_widths(table, widths[:len(active_cols)])
                
            for tx in journal_data:
                row = table.add_row().cells
                for i, col_name in enumerate(active_cols):
                    key = available_cols.get(col_name); val = tx.get(key, "")
                    if col_name == "Date": row[i].text = format_dt(str(val))
                    elif col_name == "Amount": row[i].text = format_curr(val)
                    elif col_name == "Phase": row[i].text = str(val or "Whole")
                    else: row[i].text = str(val or "")
                    row[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            apply_studio_table_styles(table)

    # LEDGER SECTION (Efficient ID-Based)
    show_ledger = (is_full_report or report_type == "Ledger")
    if sections and is_full_report:
        show_ledger = sections.get("ledger", True)

    if show_ledger:
        document.add_paragraph()
        h = document.add_paragraph()
        run = h.add_run('GENERAL LEDGER'); run.bold = True; run.font.size = Pt(16)
        
        # Load categories to resolve Codes
        all_categories = db_ref.child("categories").get() or {}
        cat_list = [c for c in all_categories.values() if isinstance(c, dict)]
        
        # Standard accounts
        standard = [{"name": "Cash", "code": 1001}, {"name": "Bank", "code": 1002}]
        cat_list = standard + cat_list

        ledger_targets = ledger_accounts if ledger_accounts else [c["name"] for c in cat_list]

        found_any = False
        for target_name in ledger_targets:
            # Resolve Code
            cat_obj = next((c for c in cat_list if c["name"] == target_name), None)
            if not cat_obj: continue
            
            # Fetch Ledger via ID-First Service (Fast!)
            entries = get_ledger_page(project_id, cat_obj["code"], phase_id=phase_id)
            if not entries: continue
            found_any = True

            if sort_order == "Descending": entries.reverse()

            p = document.add_paragraph()
            run = p.add_run(f'ACCOUNT: {target_name.upper()} (Code: {cat_obj["code"]})')
            run.bold = True; run.font.size = Pt(12)
            
            active_cols = columns.get("ledger") if (columns and columns.get("ledger")) else ['Date', 'Phase', 'Debit', 'Credit', 'Running Balance']
            table = document.add_table(rows=1, cols=len(active_cols))
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, c_name in enumerate(active_cols):
                hdr[i].text = c_name
                hdr[i].paragraphs[0].runs[0].bold = True
                hdr[i].paragraphs[0].runs[0].font.size = Pt(9.5)

            set_column_widths(table, [Inches(1.2)] * len(active_cols))
            
            for entry in entries:
                row = table.add_row().cells
                for i, c_name in enumerate(active_cols):
                    if c_name == 'Date': row[i].text = format_dt(entry["date"])
                    elif c_name == 'Phase': row[i].text = str(entry.get("phase_name") or "Whole")
                    elif c_name == 'Debit': row[i].text = format_curr(entry['amount']) if entry["entry_type"] == "Debit" else "-"
                    elif c_name == 'Credit': row[i].text = format_curr(entry['amount']) if entry["entry_type"] == "Credit" else "-"
                    elif c_name == 'Running Balance': row[i].text = format_curr(entry.get('running_balance', 0))
                    row[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            apply_studio_table_styles(table)
        
        if not found_any:
            document.add_paragraph("No ledger entries found for selected filters.")

    # TRIAL BALANCE SECTION
    show_tb = (is_full_report or report_type == "Trial Balance")
    if sections and is_full_report:
        show_tb = sections.get("trialBalance", True)

    if show_tb:
        document.add_paragraph()
        h = document.add_paragraph()
        run = h.add_run('TRIAL BALANCE'); run.bold = True; run.font.size = Pt(16)
        
        accounts_tb = trial_balance_data.get("accounts", {})
        if not accounts_tb:
            document.add_paragraph("No balances found.")
        else:
            table = document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Account Name', 'Debit Balance', 'Credit Balance'
            for cell in hdr: 
                cell.paragraphs[0].runs[0].bold = True; cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            set_column_widths(table, [Inches(3.27), Inches(1.7), Inches(1.7)])
            
            for acc, adata in accounts_tb.items():
                row = table.add_row().cells
                row[0].text = acc
                bal = adata.get("balance", 0.0)
                if bal >= 0:
                    row[1].text = format_curr(bal); row[2].text = "-"
                else:
                    row[1].text = "-"; row[2].text = format_curr(abs(bal))
                for cell in row: cell.paragraphs[0].runs[0].font.size = Pt(10)
                
            footer = table.add_row().cells
            footer[0].text = "TOTAL"; footer[0].paragraphs[0].runs[0].bold = True
            footer[1].text = format_curr(trial_balance_data['totals']['total_debits']); footer[1].paragraphs[0].runs[0].bold = True
            footer[2].text = format_curr(trial_balance_data['totals']['total_credits']); footer[2].paragraphs[0].runs[0].bold = True
            apply_studio_table_styles(table)

    # Footer Note
    if footer_note:
        document.add_paragraph()
        note_text = document.add_paragraph()
        note_run = note_text.add_run(footer_note)
        note_run.italic = True; note_run.font.size = Pt(10); note_run.font.color.rgb = RGBColor(100, 100, 100)

    temp_dir = tempfile.gettempdir()
    safe_name = "".join(x for x in project_name if x.isalnum())
    file_path = os.path.join(temp_dir, f"{safe_name}_{report_type.replace(' ', '_')}.docx")
    document.save(file_path); return file_path
